"""
文档加载器 - 支持多种文档格式
"""
import os
from typing import List, Optional
from pathlib import Path
from loguru import logger

from vectorstore.base import Document


class DocumentLoader:
    """文档加载器 - 支持多种格式"""
    
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.md': 'markdown',
        '.pdf': 'pdf',
        '.docx': 'word',
        '.doc': 'word',
        '.xlsx': 'excel',
        '.xls': 'excel',
        '.csv': 'csv',
        '.json': 'json',
        '.html': 'html',
        '.htm': 'html',
    }
    
    @classmethod
    def load(cls, file_path: str, metadata: Optional[dict] = None) -> List[Document]:
        """
        加载单个文档
        
        Args:
            file_path: 文件路径
            metadata: 附加元数据
            
        Returns:
            Document 列表
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        ext = path.suffix.lower()
        doc_type = cls.SUPPORTED_EXTENSIONS.get(ext)
        
        if not doc_type:
            logger.warning(f"不支持的文件类型: {ext}，尝试作为文本文件读取")
            doc_type = 'text'
        
        # 准备元数据
        meta = {
            'source': str(path.absolute()),
            'filename': path.name,
            'file_type': doc_type,
            **(metadata or {})
        }
        
        # 根据类型加载
        loader_method = getattr(cls, f'_load_{doc_type}', cls._load_text)
        return loader_method(file_path, meta)
    
    @classmethod
    def load_directory(
        cls,
        directory: str,
        pattern: str = "*",
        recursive: bool = True,
        metadata: Optional[dict] = None
    ) -> List[Document]:
        """
        加载目录中的所有文档
        
        Args:
            directory: 目录路径
            pattern: 文件匹配模式
            recursive: 是否递归子目录
            metadata: 附加元数据
        """
        path = Path(directory)
        if not path.exists():
            raise FileNotFoundError(f"目录不存在: {directory}")
        
        documents = []
        
        if recursive:
            files = path.rglob(pattern)
        else:
            files = path.glob(pattern)
        
        for file_path in files:
            if file_path.is_file():
                try:
                    docs = cls.load(str(file_path), metadata)
                    documents.extend(docs)
                    logger.info(f"已加载: {file_path}")
                except Exception as e:
                    logger.error(f"加载文件失败 {file_path}: {e}")
        
        logger.info(f"共加载 {len(documents)} 个文档片段")
        return documents
    
    @classmethod
    def _load_text(cls, file_path: str, metadata: dict) -> List[Document]:
        """加载文本文件"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        return [Document(page_content=content, metadata=metadata)]
    
    @classmethod
    def _load_markdown(cls, file_path: str, metadata: dict) -> List[Document]:
        """加载 Markdown 文件"""
        # Markdown 作为文本处理，可以后续添加 Markdown 解析
        return cls._load_text(file_path, metadata)
    
    @classmethod
    def _load_pdf(cls, file_path: str, metadata: dict) -> List[Document]:
        """加载 PDF 文件"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            documents = []
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    page_meta = {
                        **metadata,
                        'page_number': i + 1,
                        'total_pages': len(reader.pages)
                    }
                    documents.append(Document(page_content=text, metadata=page_meta))
            
            return documents
        except ImportError:
            logger.error("请安装 pypdf: pip install pypdf")
            raise
    
    @classmethod
    def _load_word(cls, file_path: str, metadata: dict) -> List[Document]:
        """加载 Word 文档"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            # 也读取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    content.append(' | '.join(row_text))
            
            full_text = '\n'.join(content)
            return [Document(page_content=full_text, metadata=metadata)]
            
        except ImportError:
            logger.error("请安装 python-docx: pip install python-docx")
            raise
    
    @classmethod
    def _load_excel(cls, file_path: str, metadata: dict) -> List[Document]:
        """加载 Excel 文件"""
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(file_path, data_only=True)
            documents = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                content = []
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else '' for cell in row]
                    if any(row_text):
                        content.append(' | '.join(row_text))
                
                if content:
                    sheet_meta = {
                        **metadata,
                        'sheet_name': sheet_name
                    }
                    sheet_text = '\n'.join(content)
                    documents.append(Document(page_content=sheet_text, metadata=sheet_meta))
            
            return documents
            
        except ImportError:
            logger.error("请安装 openpyxl: pip install openpyxl")
            raise
    
    @classmethod
    def _load_csv(cls, file_path: str, metadata: dict) -> List[Document]:
        """加载 CSV 文件"""
        import csv
        
        content = []
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            for row in reader:
                content.append(' | '.join(row))
        
        full_text = '\n'.join(content)
        return [Document(page_content=full_text, metadata=metadata)]
    
    @classmethod
    def _load_json(cls, file_path: str, metadata: dict) -> List[Document]:
        """加载 JSON 文件"""
        import json
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 将 JSON 转为格式化的字符串
        content = json.dumps(data, ensure_ascii=False, indent=2)
        return [Document(page_content=content, metadata=metadata)]
    
    @classmethod
    def _load_html(cls, file_path: str, metadata: dict) -> List[Document]:
        """加载 HTML 文件"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            # 移除脚本和样式
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text(separator='\n', strip=True)
            return [Document(page_content=text, metadata=metadata)]
            
        except ImportError:
            logger.warning("未安装 beautifulsoup4，将直接读取 HTML 文本")
            return cls._load_text(file_path, metadata)


class TextSplitter:
    """文本分割器 - 将长文本分割成小块"""
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separators: List[str] = None
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", "。", "，", " ", ""]
    
    def split_text(self, text: str) -> List[str]:
        """分割文本"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            if end >= len(text):
                chunks.append(text[start:])
                break
            
            # 寻找最佳分割点
            chunk = text[start:end]
            best_split = end
            
            for sep in self.separators:
                pos = chunk.rfind(sep)
                if pos > self.chunk_size * 0.3:  # 至少保留 30% 内容
                    best_split = start + pos + len(sep)
                    break
            
            chunks.append(text[start:best_split])
            start = best_split - self.chunk_overlap
        
        return chunks
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """分割文档列表"""
        result = []
        
        for doc in documents:
            chunks = self.split_text(doc.page_content)
            
            for i, chunk in enumerate(chunks):
                chunk_meta = {
                    **doc.metadata,
                    'chunk_index': i,
                    'total_chunks': len(chunks)
                }
                result.append(Document(
                    page_content=chunk,
                    metadata=chunk_meta
                ))
        
        return result


# 便捷函数
def load_document(file_path: str, metadata: Optional[dict] = None) -> List[Document]:
    """加载单个文档"""
    return DocumentLoader.load(file_path, metadata)


def load_documents(
    directory: str,
    pattern: str = "*",
    recursive: bool = True
) -> List[Document]:
    """加载目录中的所有文档"""
    return DocumentLoader.load_directory(directory, pattern, recursive)


def split_documents(
    documents: List[Document],
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> List[Document]:
    """分割文档"""
    splitter = TextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_documents(documents)
